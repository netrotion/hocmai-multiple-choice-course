data = open('datadef.txt','r',encoding='utf-8').read()
#<span class="firstletter">
import os
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from toolkit.latexconv import conv
from toolkit.download_img import get
from docx.shared import Inches



def format_to_superscript(text):
    superscript_map = str.maketrans("0123456789", "⁰¹²³⁴⁵⁶⁷⁸⁹")
    superscript_text = text.translate(superscript_map)
    return superscript_text


def final_scan(data):
    d = []
    total = data.count('<div><img src="')
    for i in range(total):
        full_content = data[data.index('<img src="'):data.index('/>')+2]
        url = data.split('<div><img src="')[1].split('"')[0]
        filename = get(url).get_data()
        content = data[:data.index(full_content)]
        data = data.replace(full_content,'').replace(content,'')
        d += ['\n'+content+'\n',filename]
        if i == total-1:
            d += ['\n'+data]
    if d == []:
        return data.replace('</div>','')
    return d
def get_data(data):
    d = {'data':[]}
    da = ['A','B','C','D']
    s = ''
    for i in range(data.count('<span class="firstletter">')):
        answer = []
        data = data[data.index('<span class="firstletter">'):]
        data1= data
        data = data.replace('<span class="firstletter">','',1)
        question = data.split('<div class="qtext">')[1].split('<div class="ablock clearfix">')[0]
        s += '\nCâu '+str(i+1)+' :'+question.replace('\n','').replace('\t','') + '\n'
        for i in range(4):
            data1 = data1[data1.index("<span class='under_an'>"):]
            data1 = data1.replace("<span class='under_an'>",'',1)
            anws = data1.split('</span>. ')[1].split('</label>')[0]
            while anws[-1] == ' ':anws = anws[:-2]
            s += da[i] + '. ' + anws + '\n' 
            answer.append(anws)
        s += '\n'
    s = s.replace('<br />','\n').replace('<sub>','').replace('</sub>','')
    #.replace('<sup>2</sup>','²')
    
    while '<sup>' in s:
        value = s.split('<sup>')[1].split('</sup>')[0]
        s = s.replace(f'<sup>{value}</sup>',format_to_superscript(value))

    while '\n\n\n' in s:
        s = s.replace('\n\n\n','\n')
    s = s.replace('\[','$$ ').replace('\]',' $$').split('$$')
    for i in range(len(s)):
        s[i] = final_scan(s[i])
    #s = flatten_list(s)
    for i in range(len(s)):
        print('process:',str(i),':',str(len(s)),end='\r')
        if i % 2 == 0:
            if 'list' in str(type(s[i])):
                for j in range(len(s[i])):
                    if '.png' not in s[i][j]:
                         d['data'].append({'type':'content','content':s[i][j]})
                         continue
                    d['data'].append({'type':'picture','dir':s[i][j]})
                continue
            d['data'].append({'type':'content','content':s[i]})
            continue
            
        d['data'].append({'type':'picture','dir':conv(s[i]).convert()})
    open('log.txt','w',encoding='utf-8').write(str(d))
    return d

def write_to_docx(file_path='t.docx', content=get_data(data)):\
    #data : {'data':[{'type':'text','content':abcd}]}
    #data : {'data':[{'type':'picture','dir':abcd}]}
    doc = Document()
    paragraph = doc.add_paragraph()
    for i in content['data']:
        if i['type'] == 'content':
            content_to_write = i['content'].replace('</div>','').replace('<div>','').replace('\n\n','\n')
            paragraph.add_run(content_to_write)
        elif i['type'] == 'picture':
            run = paragraph.add_run()
            run.add_picture(i['dir'])#, width=Inches(auto),height=Inches(auto))
    try:
        doc.save(file_path)
    except PermissionError:
        input('Close docx tab and continue with enter!')
        doc.save(file_path)
write_to_docx()