import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
import sympy as sp


latex_string = r'_{1}'
fig, ax = plt.subplots()
ax.text(0.5, 0.5, f'${latex_string}$', size=11, ha='center', va='center')
ax.axis('off')

# Lưu biểu đồ thành hình ảnh PNG
image_path = 'math_formula1.png'
plt.savefig(image_path, transparent=True)
plt.close()

doc = Document()
doc.add_paragraph('Biểu thức toán học:')
doc.add_picture(image_path, width=Inches(1))  # Chèn hình ảnh vào tài liệu, đặt chiều rộng là 3 inch

# Lưu tài liệu Word
doc.save('math_formula.docx')
